#!/usr/bin/env python3
"""Convert HTML report to Word format using LibreOffice with proper page breaks.

Handles:
- RTL text direction
- Justified alignment
- Page breaks (no table cuts)
- Paragraph integrity
"""

import sys
import subprocess
import tempfile
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def convert_html_to_docx_libreoffice(html_path, docx_path):
    """Convert HTML to DOCX using LibreOffice with proper settings."""
    html_file = Path(html_path).resolve()
    docx_file = Path(docx_path).resolve()

    if not html_file.exists():
        print(f"Error: {html_file} not found")
        return False

    try:
        # Use LibreOffice headless conversion
        # --headless: run in headless mode
        # --convert-to: specify output format
        # --outdir: output directory
        cmd = [
            'libreoffice',
            '--headless',
            '--convert-to', 'docx',
            '--outdir', str(docx_file.parent),
            str(html_file),
        ]

        print(f"Converting: {html_file} → {docx_file}")
        print(f"Command: {' '.join(cmd)}")

        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)

        if result.returncode != 0:
            print(f"LibreOffice error: {result.stderr}")
            return False

        # LibreOffice creates file with same basename, rename if needed
        libreoffice_output = docx_file.parent / html_file.stem
        libreoffice_output = libreoffice_output.with_suffix('.docx')

        if libreoffice_output != docx_file and libreoffice_output.exists():
            libreoffice_output.rename(docx_file)

        print(f"✓ Conversion complete: {docx_file}")
        return True

    except subprocess.TimeoutExpired:
        print("Error: LibreOffice conversion timeout")
        return False
    except Exception as e:
        print(f"Error: {e}")
        return False


def enhance_docx_rtl_settings(docx_path):
    """Open existing DOCX and enhance RTL/justification settings."""
    try:
        doc = Document(docx_path)

        # Set document-level RTL
        settings = doc.settings._element
        if settings is not None:
            # Ensure bidi (right-to-left) is set
            bidi = OxmlElement('w:bidi')
            settings.insert(0, bidi)

        # Process all paragraphs for RTL and justification
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        for paragraph in doc.paragraphs:
            # Set RTL direction
            pPr = paragraph._element.get_or_add_pPr()
            bidi = pPr.find(qn('w:bidi'))
            if bidi is None:
                bidi = OxmlElement('w:bidi')
                pPr.append(bidi)

            # Set language to Hebrew
            rPr = pPr.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                pPr.append(rPr)

            lang = rPr.find(qn('w:lang'))
            if lang is None:
                lang = OxmlElement('w:lang')
                lang.set(qn('w:val'), 'he-IL')
                rPr.append(lang)

            # Set alignment to justified (if not already)
            if paragraph.alignment is None:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Process table cells for RTL
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        pPr = para._element.get_or_add_pPr()
                        bidi = pPr.find(qn('w:bidi'))
                        if bidi is None:
                            bidi = OxmlElement('w:bidi')
                            pPr.append(bidi)

        doc.save(docx_path)
        print(f"✓ Enhanced RTL settings: {docx_path}")
        return True

    except Exception as e:
        print(f"Error enhancing DOCX: {e}")
        return False


def add_page_break_before_tables(docx_path):
    """Add page breaks before tables to avoid cutting them."""
    try:
        doc = Document(docx_path)

        # Track if we should add a page break
        prev_was_para = False

        for i, element in enumerate(doc.element.body):
            # Check if this is a table
            if element.tag == qn('w:tbl'):
                # Add page break before table if there was content before
                if i > 0 and prev_was_para:
                    # Insert page break
                    para_before = doc.element.body[i - 1]
                    if para_before.tag == qn('w:p'):
                        # Add break to last paragraph
                        p = para_before
                        pPr = p.get_or_add_pPr()
                        br = OxmlElement('w:br')
                        br.set(qn('w:type'), 'page')
                        pPr.append(br)

                prev_was_para = False
            elif element.tag == qn('w:p'):
                prev_was_para = True

        doc.save(docx_path)
        print(f"✓ Added page break protection: {docx_path}")
        return True

    except Exception as e:
        print(f"Error adding page breaks: {e}")
        return False


def convert_html_to_docx(html_path, docx_path, enhance_rtl=True):
    """Main conversion function: HTML → DOCX with all enhancements."""
    print(f"\n{'='*60}")
    print(f"HTML → DOCX Conversion Pipeline")
    print(f"{'='*60}\n")

    # Step 1: Convert using LibreOffice
    print("Step 1: LibreOffice HTML→DOCX conversion")
    if not convert_html_to_docx_libreoffice(html_path, docx_path):
        return False

    # Step 2: Enhance RTL and language settings
    if enhance_rtl:
        print("\nStep 2: Enhance RTL and language settings")
        enhance_docx_rtl_settings(docx_path)

    # Step 3: Add page break protection
    print("\nStep 3: Add page break protection for tables")
    add_page_break_before_tables(docx_path)

    print(f"\n{'='*60}")
    print(f"✓ Conversion complete!")
    print(f"Output: {Path(docx_path).name}")
    print(f"Size: {Path(docx_path).stat().st_size / 1024:.1f} KB")
    print(f"{'='*60}\n")

    return True


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python html_to_docx.py <input.html> <output.docx>")
        sys.exit(1)

    html_file = sys.argv[1]
    docx_file = sys.argv[2]

    success = convert_html_to_docx(html_file, docx_file)
    sys.exit(0 if success else 1)
