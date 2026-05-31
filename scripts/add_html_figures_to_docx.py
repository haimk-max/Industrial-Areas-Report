#!/usr/bin/env python3
"""Add charts and maps from HTML to Word document.

Extracts base64-encoded images from HTML and embeds them in Word document.
Processes SVG figures and inserts them as images at appropriate locations.
"""

import sys
import re
import base64
import io
from pathlib import Path
from PIL import Image
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def extract_figures_from_html(html_path):
    """Extract embedded figures (base64 images) from HTML."""
    with open(html_path, 'r', encoding='utf-8') as f:
        html = f.read()

    figures = []

    # Look for <img> tags with base64 data URIs
    img_pattern = r'<img[^>]*src="data:image/([^;]+);base64,([^"]+)"[^>]*alt="([^"]*)"[^>]*>'
    for match in re.finditer(img_pattern, html):
        img_type, b64_data, alt_text = match.groups()
        try:
            img_bytes = base64.b64decode(b64_data)
            figures.append({
                'type': img_type,
                'data': img_bytes,
                'alt': alt_text,
                'format': 'PNG' if img_type == 'png' else 'SVG'
            })
        except Exception as e:
            print(f"Warning: Failed to decode image: {e}", file=sys.stderr)

    # Also look for SVG inline (embedded in figure tags or direct)
    svg_pattern = r'(<svg[^>]*>.*?</svg>)'
    for match in re.finditer(svg_pattern, html, re.DOTALL):
        svg_text = match.group(1)
        # Try to get caption nearby
        caption_pattern = r'<figcaption>([^<]+)</figcaption>'
        caption_match = re.search(caption_pattern, html[max(0, match.start()-200):match.end()])
        caption = caption_match.group(1) if caption_match else "תרשים"

        figures.append({
            'type': 'svg',
            'data': svg_text.encode('utf-8'),
            'alt': caption,
            'format': 'SVG'
        })

    return figures


def svg_to_png(svg_bytes):
    """Convert SVG to PNG using cairosvg or PIL."""
    try:
        import cairosvg
        png_bytes = io.BytesIO()
        cairosvg.svg2png(bytestring=svg_bytes, write_to=png_bytes)
        png_bytes.seek(0)
        return png_bytes
    except ImportError:
        # Fallback: try using PIL with SVG support
        try:
            from PIL import Image
            img = Image.open(io.BytesIO(svg_bytes))
            png_bytes = io.BytesIO()
            img.save(png_bytes, format='PNG')
            png_bytes.seek(0)
            return png_bytes
        except Exception as e:
            print(f"Warning: Could not convert SVG to PNG: {e}", file=sys.stderr)
            return None


def add_figures_to_docx(docx_path, html_path, output_docx_path=None):
    """Add figures from HTML to Word document."""
    if output_docx_path is None:
        output_docx_path = docx_path

    # Extract figures from HTML
    print(f"Extracting figures from: {html_path}")
    figures = extract_figures_from_html(html_path)
    print(f"Found {len(figures)} figures")

    if not figures:
        print("No figures to add")
        return

    # Open existing Word document
    doc = Document(docx_path)

    # Try to insert figures after specific markers or at section boundaries
    # For now, add figures at the end with captions
    doc.add_heading('תרשימים ותמונות', level=1)

    added_count = 0
    for fig_idx, fig in enumerate(figures):
        try:
            if fig['type'] == 'svg':
                # Convert SVG to PNG
                png_data = svg_to_png(fig['data'])
                if png_data:
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run()
                    run.add_picture(png_data, width=Inches(5.5))
                    added_count += 1
            else:
                # Direct image (already PNG/JPG)
                img_bytes = io.BytesIO(fig['data'])
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(img_bytes, width=Inches(5.5))
                added_count += 1

            # Add caption
            caption_para = doc.add_paragraph(fig['alt'])
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in caption_para.runs:
                run.font.size = 9
                run.font.italic = True
            caption_para.paragraph_format.space_before = 6
            caption_para.paragraph_format.space_after = 12

        except Exception as e:
            print(f"Warning: Failed to add figure {fig_idx}: {e}", file=sys.stderr)

    print(f"Added {added_count}/{len(figures)} figures to Word document")

    # Save modified document
    doc.save(output_docx_path)
    print(f"✓ Saved: {output_docx_path}")


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python add_html_figures_to_docx.py <input.docx> <input.html> [output.docx]")
        sys.exit(1)

    docx_file = Path(sys.argv[1])
    html_file = Path(sys.argv[2])
    output_file = Path(sys.argv[3]) if len(sys.argv) > 3 else None

    if not docx_file.exists() or not html_file.exists():
        print(f"Error: Input files not found")
        sys.exit(1)

    add_figures_to_docx(str(docx_file), str(html_file), str(output_file) if output_file else None)
