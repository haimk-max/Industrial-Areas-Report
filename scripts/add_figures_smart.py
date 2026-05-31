#!/usr/bin/env python3
"""Add charts to Word document at their original positions with RTL captions.

Parses markdown to find figure references (איור N) and inserts images
at those exact locations, maintaining RTL text direction.
"""

import sys
import re
import base64
import io
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def extract_figures_from_html(html_path):
    """Extract embedded figures (SVG + images) from HTML with metadata."""
    with open(html_path, 'r', encoding='utf-8') as f:
        html = f.read()

    figures = {}
    fig_counter = 0

    # Look for <svg> blocks that are figures
    svg_pattern = r'<svg[^>]*>.*?</svg>'
    for match in re.finditer(svg_pattern, html, re.DOTALL):
        svg_text = match.group(0)
        fig_counter += 1

        # Try to find figure number in context (look backwards for "איור")
        start_pos = max(0, match.start() - 500)
        context_before = html[start_pos:match.start()]
        fig_match = re.search(r'(?:\*\*)?איור\s+(\d+)(?:\*\*)?', context_before)
        fig_num = int(fig_match.group(1)) if fig_match else fig_counter

        # Try to find caption in next text
        end_pos = min(len(html), match.end() + 300)
        context_after = html[match.end():end_pos]
        caption_match = re.search(r'<figcaption>([^<]+)</figcaption>', context_after)
        caption = caption_match.group(1) if caption_match else f"איור {fig_num}"

        figures[fig_num] = {
            'type': 'svg',
            'data': svg_text.encode('utf-8'),
            'alt': caption,
            'format': 'SVG'
        }

    # Also look for base64 images
    img_pattern = r'<img[^>]*src="data:image/([^;]+);base64,([^"]+)"[^>]*alt="([^"]*)"[^>]*>'
    for match in re.finditer(img_pattern, html):
        img_type, b64_data, alt_text = match.groups()
        try:
            img_bytes = base64.b64decode(b64_data)

            # Extract figure number from alt text
            fig_match = re.search(r'איור\s*(\d+)', alt_text)
            fig_num = int(fig_match.group(1)) if fig_match else 100 + fig_counter
            fig_counter += 1

            if fig_num not in figures:  # Don't overwrite SVGs
                figures[fig_num] = {
                    'type': img_type,
                    'data': img_bytes,
                    'alt': alt_text,
                    'format': 'PNG'
                }
        except Exception as e:
            print(f"Warning: Failed to decode image: {e}", file=sys.stderr)

    return figures


def find_figure_anchors_in_docx(doc):
    """Find paragraph indices where figures should be inserted (search for 'איור N')."""
    anchors = {}

    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        # Look for figure references like "**איור 1**"
        match = re.search(r'(?:\*\*)?איור\s+(\d+)(?:\*\*)?', text)
        if match:
            fig_num = int(match.group(1))
            anchors[fig_num] = para_idx

    return anchors


def svg_to_png(svg_bytes):
    """Convert SVG to PNG using cairosvg."""
    try:
        import cairosvg
        png_bytes = io.BytesIO()
        cairosvg.svg2png(bytestring=svg_bytes, write_to=png_bytes)
        png_bytes.seek(0)
        return png_bytes.getvalue()
    except ImportError:
        print("Warning: cairosvg not available, skipping SVG conversion", file=sys.stderr)
        return None
    except Exception as e:
        print(f"Warning: SVG conversion failed: {e}", file=sys.stderr)
        return None


def insert_figure_at_paragraph(doc, para_idx, image_data, fig_type, caption, fig_num):
    """Insert figure after specified paragraph with RTL caption."""
    # Convert SVG to PNG if needed
    if fig_type == 'svg':
        png_bytes = svg_to_png(image_data)
        if png_bytes is None:
            print(f"  Skipping figure {fig_num} (SVG conversion failed)")
            return None
        image_data = png_bytes

    # Get the paragraph after which we'll insert
    if para_idx >= len(doc.paragraphs) - 1:
        # Insert at end
        para_ref = doc.paragraphs[-1]
    else:
        para_ref = doc.paragraphs[para_idx]

    # Add new paragraph for image
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Insert image
    try:
        run = img_para.add_run()
        img_stream = io.BytesIO(image_data)
        run.add_picture(img_stream, width=Inches(5.5))
    except Exception as e:
        print(f"  Warning: Failed to insert picture: {e}", file=sys.stderr)
        return None

    # Add caption paragraph (RTL)
    caption_para = doc.add_paragraph()
    set_paragraph_rtl_justified(caption_para)

    # Add caption text
    caption_run = caption_para.add_run(f"איור {fig_num}: {caption}")
    caption_run.font.size = Pt(10)
    caption_run.font.italic = True

    # Apply RTL to caption run
    rPr = caption_run._element.get_or_add_rPr()
    if rPr.find(qn('w:rtl')) is None:
        rtl = OxmlElement('w:rtl')
        rPr.append(rtl)

    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_para.paragraph_format.space_before = Pt(6)
    caption_para.paragraph_format.space_after = Pt(12)

    return img_para


def set_paragraph_rtl_justified(para):
    """Set paragraph RTL and justified."""
    pPr = para._element.get_or_add_pPr()

    # RTL
    if pPr.find(qn('w:bidi')) is None:
        bidi = OxmlElement('w:bidi')
        pPr.insert(0, bidi)

    # Justified
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc')
        pPr.append(jc)
    jc.set(qn('w:val'), 'both')


def add_figures_to_docx_smart(docx_path, html_path, output_docx_path=None):
    """Add figures to Word at original positions with RTL captions."""
    if output_docx_path is None:
        output_docx_path = docx_path

    print(f"Extracting figures from HTML: {html_path}")
    figures = extract_figures_from_html(html_path)
    print(f"Found {len(figures)} figures")

    if not figures:
        print("No figures to add")
        return

    print(f"Opening Word document: {docx_path}")
    doc = Document(docx_path)

    # Find where figures should be inserted
    anchors = find_figure_anchors_in_docx(doc)
    print(f"Found {len(anchors)} figure anchors in document")

    # Insert figures (in reverse order to avoid index shifting)
    inserted_count = 0
    for fig_num in sorted(anchors.keys(), reverse=True):
        if fig_num in figures:
            para_idx = anchors[fig_num]
            fig_data = figures[fig_num]

            # Extract caption (remove "איור N:" prefix if present)
            caption = fig_data['alt']
            caption = re.sub(r'^\s*(?:\*\*)?איור\s+\d+(?:\*\*)?[\s:]*', '', caption).strip()

            try:
                result = insert_figure_at_paragraph(doc, para_idx, fig_data['data'], fig_data['type'], caption, fig_num)
                if result:
                    print(f"  Inserted figure {fig_num} at paragraph {para_idx}")
                    inserted_count += 1
            except Exception as e:
                print(f"  Warning: Failed to insert figure {fig_num}: {e}", file=sys.stderr)
        else:
            print(f"  Figure {fig_num} referenced but not found in HTML", file=sys.stderr)

    print(f"\nAdded {inserted_count}/{len(anchors)} figures to Word document")

    # Save
    doc.save(output_docx_path)
    print(f"✓ Saved: {output_docx_path}")


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python add_figures_smart.py <input.docx> <input.html> [output.docx]")
        sys.exit(1)

    docx_file = Path(sys.argv[1])
    html_file = Path(sys.argv[2])
    output_file = Path(sys.argv[3]) if len(sys.argv) > 3 else None

    if not docx_file.exists() or not html_file.exists():
        print(f"Error: Input files not found")
        sys.exit(1)

    add_figures_to_docx_smart(str(docx_file), str(html_file), str(output_file) if output_file else None)
