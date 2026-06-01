#!/usr/bin/env python3
"""Embed SVG figures from HTML into Word document at original positions.

Reads HOLON_REPORT_V5.html for SVG figure blocks.
Extracts figure numbers, SVG content, and captions.
Converts SVG→PNG using cairosvg.
Inserts PNG images at **איור N** anchors in Word document.
"""

import sys
import re
import io
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPO_ROOT = Path(__file__).resolve().parent.parent


def extract_svg_figures(html_path):
    """Extract all SVG figures from HTML with figure numbers."""
    with open(html_path, encoding='utf-8') as f:
        html = f.read()

    figures = {}

    # Find SVG blocks (each SVG is a complete figure)
    svg_pattern = r'<svg[^>]*>.*?</svg>'
    for match in re.finditer(svg_pattern, html, re.DOTALL):
        svg_text = match.group(0)

        # Try to find figure number from context (look backwards for איור N)
        start = max(0, match.start() - 300)
        context = html[start:match.start()]
        fig_match = re.search(r'(?:\*\*)?איור\s+(\d+)(?:\*\*)?', context)

        if fig_match:
            fig_num = int(fig_match.group(1))
            figures[fig_num] = {
                'svg': svg_text,
                'type': 'svg'
            }

    return figures


def svg_to_png(svg_bytes):
    """Convert SVG bytes to PNG bytes using cairosvg."""
    try:
        import cairosvg
        png_io = io.BytesIO()
        cairosvg.svg2png(bytestring=svg_bytes, write_to=png_io)
        png_io.seek(0)
        return png_io.getvalue()
    except Exception as e:
        print(f"Warning: SVG conversion failed: {e}", file=sys.stderr)
        return None


def find_figure_anchors(doc):
    """Find paragraph indices with **איור N** markers."""
    anchors = {}
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        match = re.search(r'(?:\*\*)?איור\s+(\d+)(?:\*\*)?', text)
        if match:
            fig_num = int(match.group(1))
            anchors[fig_num] = idx
    return anchors


def set_paragraph_rtl_justified(para):
    """Configure paragraph for RTL Hebrew with justified alignment."""
    pPr = para._element.get_or_add_pPr()

    # RTL flag
    if pPr.find(qn('w:bidi')) is None:
        bidi = OxmlElement('w:bidi')
        pPr.insert(0, bidi)

    # Justified alignment
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc')
        pPr.append(jc)
    jc.set(qn('w:val'), 'both')


def embed_figures(docx_path, html_path, output_path=None):
    """Embed extracted figures into Word document."""
    if output_path is None:
        output_path = docx_path

    print(f"Extracting figures from HTML: {html_path}")
    figures = extract_svg_figures(html_path)
    print(f"  Found {len(figures)} SVG figures")

    if not figures:
        print("No figures to embed, saving original document")
        doc = Document(docx_path)
        doc.save(output_path)
        return

    print(f"Loading Word document: {docx_path}")
    doc = Document(docx_path)

    anchors = find_figure_anchors(doc)
    print(f"  Found {len(anchors)} figure anchors in document")

    # Insert figures (reverse order to avoid index shifting)
    inserted = 0
    for fig_num in sorted(anchors.keys(), reverse=True):
        if fig_num not in figures:
            print(f"  Warning: figure {fig_num} referenced but not found in HTML")
            continue

        para_idx = anchors[fig_num]
        svg_text = figures[fig_num]['svg']

        # Convert SVG to PNG
        png_bytes = svg_to_png(svg_text.encode('utf-8'))
        if not png_bytes:
            print(f"  Skipping figure {fig_num} (conversion failed)")
            continue

        # Insert image paragraph after the anchor paragraph
        if para_idx < len(doc.paragraphs) - 1:
            insert_para = doc.paragraphs[para_idx]._element
            insert_para = insert_para.addnext(doc.add_paragraph()._element)
        else:
            insert_para = doc.add_paragraph()

        img_para = insert_para if hasattr(insert_para, 'addnext') else doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        try:
            run = img_para.add_run()
            img_stream = io.BytesIO(png_bytes)
            run.add_picture(img_stream, width=Inches(5.5))
        except Exception as e:
            print(f"  Warning: Failed to insert image {fig_num}: {e}")
            continue

        # Add caption paragraph (RTL, centered)
        caption_para = doc.add_paragraph()
        set_paragraph_rtl_justified(caption_para)
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        caption_text = f"איור {fig_num}: תרשים ניתוח זיהום"
        caption_run = caption_para.add_run(caption_text)
        caption_run.font.size = Pt(10)
        caption_run.font.italic = True

        # RTL on run
        rPr = caption_run._element.get_or_add_rPr()
        if rPr.find(qn('w:rtl')) is None:
            rtl = OxmlElement('w:rtl')
            rPr.append(rtl)

        caption_para.paragraph_format.space_before = Pt(6)
        caption_para.paragraph_format.space_after = Pt(12)

        print(f"  ✓ Inserted figure {fig_num}")
        inserted += 1

    print(f"\nInserted {inserted}/{len(anchors)} figures")

    doc.save(output_path)
    print(f"✓ Saved: {output_path}")


def main():
    parser_help = "Embed SVG figures from HTML into Word document"
    if len(sys.argv) < 3:
        print(f"Usage: python {sys.argv[0]} <input.docx> <input.html> [output.docx]")
        print(parser_help)
        sys.exit(1)

    docx_file = Path(sys.argv[1])
    html_file = Path(sys.argv[2])
    output_file = Path(sys.argv[3]) if len(sys.argv) > 3 else docx_file

    if not docx_file.exists() or not html_file.exists():
        print(f"Error: Input files not found")
        sys.exit(1)

    embed_figures(str(docx_file), str(html_file), str(output_file))


if __name__ == '__main__':
    main()
