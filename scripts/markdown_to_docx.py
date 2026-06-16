#!/usr/bin/env python3
"""Convert Holon V5 Markdown report to Word (.docx) format with full RTL Hebrew support.

Properly sets:
- Document-level RTL (w:bidi)
- Paragraph-level RTL direction (w:bidi)
- Paragraph-level justified alignment (w:jc=both)
- Run-level RTL direction (w:rtl) — REQUIRED for actual RTL text rendering
- Hebrew language with bidi attribute (w:lang w:val=he-IL w:bidi=he-IL)
- Track changes enabled
"""

import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


def make_rtl_rPr():
    """Create rPr element with RTL and Hebrew language settings."""
    rPr = OxmlElement('w:rPr')

    # Set RTL flag — this is what makes Word render text RTL
    rtl = OxmlElement('w:rtl')
    rPr.append(rtl)

    # Set bidi (CS) font matching Latin font
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    rFonts.set(qn('w:cs'), 'Arial')  # Complex Script font for Hebrew
    rPr.append(rFonts)

    # Set language with both w:val and w:bidi
    lang = OxmlElement('w:lang')
    lang.set(qn('w:val'), 'en-US')
    lang.set(qn('w:bidi'), 'he-IL')  # Bidi language is Hebrew
    rPr.append(lang)

    return rPr


def apply_rtl_to_run(run):
    """Apply RTL formatting to an existing run."""
    rPr = run._element.get_or_add_rPr()

    # Add w:rtl if not present
    if rPr.find(qn('w:rtl')) is None:
        rtl = OxmlElement('w:rtl')
        rPr.append(rtl)

    # Set CS font
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:cs'), 'Arial')

    # Set bidi language
    lang = rPr.find(qn('w:lang'))
    if lang is None:
        lang = OxmlElement('w:lang')
        rPr.append(lang)
    lang.set(qn('w:bidi'), 'he-IL')


def set_paragraph_rtl_justified(paragraph):
    """Set paragraph to RTL with justified alignment."""
    pPr = paragraph._element.get_or_add_pPr()

    # Set paragraph direction to RTL (must be first child of pPr)
    if pPr.find(qn('w:bidi')) is None:
        bidi = OxmlElement('w:bidi')
        # Insert at beginning
        pPr.insert(0, bidi)

    # Set alignment to justified (both sides)
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc')
        pPr.append(jc)
    jc.set(qn('w:val'), 'both')

    # Set paragraph-level rPr for default run properties
    pPr_rPr = pPr.find(qn('w:rPr'))
    if pPr_rPr is None:
        pPr_rPr = OxmlElement('w:rPr')
        pPr.append(pPr_rPr)

    if pPr_rPr.find(qn('w:rtl')) is None:
        rtl = OxmlElement('w:rtl')
        pPr_rPr.append(rtl)

    # Apply RTL to all existing runs in this paragraph
    for run in paragraph.runs:
        apply_rtl_to_run(run)


def set_document_rtl(doc):
    """Set document-wide RTL settings."""
    settings = doc.settings._element

    # Enable trackRevisions
    if settings.find(qn('w:trackRevisions')) is None:
        track = OxmlElement('w:trackRevisions')
        settings.insert(0, track)

    # Set default tab stop
    if settings.find(qn('w:defaultTabStop')) is None:
        tab = OxmlElement('w:defaultTabStop')
        tab.set(qn('w:val'), '720')
        settings.append(tab)

    # Set theme fonts for Complex Script (Hebrew)
    # This affects the default rendering


def set_default_styles_rtl(doc):
    """Configure the Normal style to be RTL and justified."""
    style = doc.styles['Normal']

    # Set font
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Get the style element
    style_element = style._element

    # Get or create pPr
    pPr = style_element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        style_element.append(pPr)

    # Add bidi to pPr
    if pPr.find(qn('w:bidi')) is None:
        bidi = OxmlElement('w:bidi')
        pPr.insert(0, bidi)

    # Add justified alignment
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc')
        pPr.append(jc)
    jc.set(qn('w:val'), 'both')

    # Get or create rPr
    rPr = style_element.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        style_element.append(rPr)

    # Add rtl to rPr
    if rPr.find(qn('w:rtl')) is None:
        rtl = OxmlElement('w:rtl')
        rPr.append(rtl)

    # Set CS font
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    rFonts.set(qn('w:cs'), 'Arial')

    # Set language
    lang = rPr.find(qn('w:lang'))
    if lang is None:
        lang = OxmlElement('w:lang')
        rPr.append(lang)
    lang.set(qn('w:val'), 'en-US')
    lang.set(qn('w:bidi'), 'he-IL')


def parse_markdown(content):
    """Parse markdown content into blocks."""
    blocks = []
    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]

        if not line.strip():
            i += 1
            continue

        # Headings
        match = re.match(r'^(#{1,6})\s+(.*)', line)
        if match:
            level = len(match.group(1))
            text = match.group(2)
            blocks.append(('heading', level, text))
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^-{3,}$', line.strip()):
            blocks.append(('hr',))
            i += 1
            continue

        # Tables
        if line.strip().startswith('|'):
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            blocks.append(('table', table_lines))
            continue

        # Blockquote
        if line.strip().startswith('>'):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                quote_lines.append(lines[i].strip()[1:].strip())
                i += 1
            blocks.append(('quote', ' '.join(quote_lines)))
            continue

        # Lists
        if re.match(r'^\s*[-*]\s+', line):
            list_items = []
            while i < len(lines) and re.match(r'^\s*[-*]\s+', lines[i]):
                m = re.match(r'^\s*[-*]\s+(.*)', lines[i])
                list_items.append(m.group(1))
                i += 1
            blocks.append(('list', list_items))
            continue

        # Numbered list
        if re.match(r'^\s*\d+\.\s+', line):
            list_items = []
            while i < len(lines) and re.match(r'^\s*\d+\.\s+', lines[i]):
                m = re.match(r'^\s*\d+\.\s+(.*)', lines[i])
                list_items.append(m.group(1))
                i += 1
            blocks.append(('numbered_list', list_items))
            continue

        # Paragraphs
        para_lines = []
        while i < len(lines):
            line = lines[i]
            if (not line.strip() or
                re.match(r'^#{1,6}\s', line) or
                line.strip().startswith('|') or
                line.strip().startswith('>') or
                re.match(r'^-{3,}$', line.strip()) or
                re.match(r'^\s*[-*]\s+', line) or
                re.match(r'^\s*\d+\.\s+', line)):
                break
            para_lines.append(line)
            i += 1

        if para_lines:
            text = ' '.join(para_lines).strip()
            if text:
                blocks.append(('paragraph', text))

    return blocks


def add_paragraph_with_formatting(doc, text, style=None):
    """Add a paragraph with inline formatting (bold, italic, code)."""
    if style:
        para = doc.add_paragraph(style=style)
    else:
        para = doc.add_paragraph()

    # Parse inline formatting
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)

    for part in parts:
        if not part:
            continue

        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = para.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = para.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.color.rgb = RGBColor(128, 0, 0)
        else:
            para.add_run(part)

    # Apply RTL+justified to paragraph and all runs
    set_paragraph_rtl_justified(para)
    return para


def parse_table(table_lines):
    """Parse markdown table into rows."""
    rows = []
    for line in table_lines:
        stripped = line.strip()
        if stripped.startswith('|') and '---' not in stripped:
            cells = [cell.strip() for cell in stripped.split('|')[1:-1]]
            rows.append(cells)
    return rows


def markdown_to_docx(md_path, docx_path):
    """Convert markdown file to Word document with full RTL Hebrew support."""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # Configure document-wide RTL
    set_document_rtl(doc)
    set_default_styles_rtl(doc)

    # Parse and convert
    blocks = parse_markdown(content)

    for block in blocks:
        if block[0] == 'heading':
            _, level, text = block
            heading = doc.add_heading(text, level=level)
            set_paragraph_rtl_justified(heading)

        elif block[0] == 'paragraph':
            _, text = block
            add_paragraph_with_formatting(doc, text)

        elif block[0] == 'quote':
            _, text = block
            para = add_paragraph_with_formatting(doc, text, style='Intense Quote')

        elif block[0] == 'hr':
            para = doc.add_paragraph('─' * 50)
            set_paragraph_rtl_justified(para)

        elif block[0] == 'table':
            _, table_lines = block
            rows = parse_table(table_lines)

            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = 'Light Grid Accent 1'

                # Set table direction to RTL
                tbl = table._element
                tblPr = tbl.find(qn('w:tblPr'))
                if tblPr is not None:
                    if tblPr.find(qn('w:bidiVisual')) is None:
                        bidi_visual = OxmlElement('w:bidiVisual')
                        tblPr.append(bidi_visual)

                for row_idx, row_data in enumerate(rows):
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < len(table.rows[row_idx].cells):
                            cell = table.rows[row_idx].cells[col_idx]
                            para = cell.paragraphs[0]
                            # Clear default run and add formatted
                            para.text = ''
                            # Parse inline formatting in cell
                            parts = re.split(r'(\*\*[^*]+\*\*)', cell_text)
                            for part in parts:
                                if part.startswith('**') and part.endswith('**'):
                                    run = para.add_run(part[2:-2])
                                    run.bold = True
                                elif part:
                                    para.add_run(part)
                            set_paragraph_rtl_justified(para)

                            if row_idx == 0:
                                # Header row: bold all runs
                                for run in para.runs:
                                    run.bold = True

        elif block[0] == 'list':
            _, items = block
            for item in items:
                para = add_paragraph_with_formatting(doc, item, style='List Bullet')

        elif block[0] == 'numbered_list':
            _, items = block
            for item in items:
                para = add_paragraph_with_formatting(doc, item, style='List Number')

    # Save
    doc.save(docx_path)

    # Verify the output
    verify_doc = Document(docx_path)
    n_paras = len(verify_doc.paragraphs)
    n_tables = len(verify_doc.tables)

    # Check RTL on first non-empty paragraph
    rtl_count = 0
    jc_count = 0
    run_rtl_count = 0
    for p in verify_doc.paragraphs:
        pPr = p._element.find(qn('w:pPr'))
        if pPr is not None:
            if pPr.find(qn('w:bidi')) is not None:
                rtl_count += 1
            jc = pPr.find(qn('w:jc'))
            if jc is not None and jc.get(qn('w:val')) == 'both':
                jc_count += 1
        for run in p.runs:
            rPr = run._element.find(qn('w:rPr'))
            if rPr is not None and rPr.find(qn('w:rtl')) is not None:
                run_rtl_count += 1
                break

    print(f"✓ Word document created: {docx_path}")
    print(f"  - Total paragraphs: {n_paras}")
    print(f"  - Total tables: {n_tables}")
    print(f"  - Paragraphs with RTL (w:bidi): {rtl_count}/{n_paras}")
    print(f"  - Paragraphs with justified (w:jc=both): {jc_count}/{n_paras}")
    print(f"  - Paragraphs with at least one RTL run: {run_rtl_count}/{n_paras}")
    print(f"  - Track changes: enabled")
    print(f"  - Hebrew language (he-IL bidi): set")


if __name__ == '__main__':
    if len(sys.argv) != 3:
        print("Usage: python markdown_to_docx.py <input.md> <output.docx>")
        sys.exit(1)

    md_file = Path(sys.argv[1])
    docx_file = Path(sys.argv[2])

    if not md_file.exists():
        print(f"Error: {md_file} not found")
        sys.exit(1)

    markdown_to_docx(str(md_file), str(docx_file))
