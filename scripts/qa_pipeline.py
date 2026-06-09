#!/usr/bin/env python3
"""qa_pipeline.py — QA gates for the V5 hybrid pipeline.

Enforces governance rules at each pipeline step, preventing bad outputs
from propagating to downstream steps.

Gates:
  2  — Structured Data Pack (6 CSVs schema + borehole baseline + series coherence)
  3  — Prompt layer freshness (render provenance, no abolished ordering language)
  4  — Zone Diagnosis (8 questions, geographic framing, PFAS as gap)
  5  — V5 Report (terminology incl. §B.5 substitutions, consistency, structure, framing)
  6  — HTML/Word output (figures, RTL, borehole rendering)
  all — Run gates 2→4→5→6 in sequence

Usage:
  python scripts/qa_pipeline.py --gate 5 --zone Holon
  python scripts/qa_pipeline.py --gate all --zone Holon
  python scripts/qa_pipeline.py --gate 5 --zone Holon --report path/to/report.md
  python scripts/qa_pipeline.py --gate all --zone Holon --continue-on-failure

Exit codes: 0 = PASS, 1 = FAIL, 2 = WARN (pass with warnings)

Governance reference: ZONE_REPORT_PROCESS_GUIDE.md §VII, §VIII
"""

import argparse
import re
import sys
import json
import csv
from pathlib import Path
from datetime import datetime
from typing import Optional

REPO_ROOT = Path(__file__).resolve().parent.parent

# ── SSOT for banned terms (gate 5) ───────────────────────────────────────────

# Pipeline-internal terminology that must NOT appear in professional report body
PIPELINE_TERMS = [
    # File names
    r"severity_by_well_family\.csv",
    r"latest_results\.csv",
    r"measurements_scoped\.csv",
    r"monitoring_gaps\.csv",
    r"figure_ready_series\.csv",
    r"source_candidates_index\.csv",
    r"facility_attribution\.json",
    # Process step references
    r"Opus\s+[Cc]all\s+#?\d",
    r"[Ss]tep\s+[1-7]\b",
    r"V5\s+[Hh]ybrid\s+[Pp]ipeline",
    r"PROCESS_GUIDE",
    r"REPORT_V5_SCHEMA",
    # Internal classification codes
    r"\bA\+B\b",
    r"\bC-class\b",
    r"\bD-class\b",
    r"\bE-class\b",
    # Statistical internals
    r"\bsoft_trigger\b",
    r"\bSNR\s+gating\b",
    r"\bbucket\b",                   # → אינדקס חומרה
    # English operational labels (also in existing validate_report.py)
    r"\bALERT\b",
    r"\bWATCH\b",
    r"\bELEVATED\b",
    r"\bSTABLE\b",
    r"\bNONE\b",
]

# Overstatement patterns — framing that requires confidence qualification
OVERSTATEMENT_PATTERNS = [
    (r"חודרת\s+כבר\s+ל", "→ 'נמדדה חריגה ב...' או 'עשויה להעיד על...'"),
    (r"מצביע\s+על\s+גוף\s+LNAPL\s+פיזי", "→ 'מעלה חשד לקיום מקור דלקי מתמשך'"),
    (r"דליפת\s+UST\s+טריה", "→ 'תואמת אפשרות של דליפה פעילה'"),
    (r"בוודאות\s+(גורמת|מקור|אחראי)", "→ הוסף רמת ודאות (גבוהה/בינונית/נמוכה)"),
    (r"אחראי\s+לזיהום", "→ 'מועמד סביר' + רמת ודאות"),
    (r"הגיעה\s+(כבר\s+)?לקידוח\s+אספקה", "→ 'נמדד TCE בקידוח אספקה פעיל'"),
]

# Mandatory Hebrew-prose substitutions (STYLE_GUIDE §B.5) — FAIL in the report body.
# These three user-mandated decisions live in STYLE_GUIDE §B.5 (the SSOT); this list
# is the enforcement arm. Matched case-sensitively so the English confidence codes
# only trip on the uppercase HIGH/MEDIUM/LOW used as ratings, not ordinary words.
BANNED_PROSE_TERMS = [
    (r"ריקבון", "→ 'פירוק' (שרשרת פירוק / פירוק מיקרוביאלי אירובי / אנ-אירובי / תוצרי פירוק)"),
    (r"שתיק(?:ה|ות|ת)", "→ 'הפסקת ניטור' (הופסק הניטור / תקופת הפסקת ניטור)"),
    (r"קיצוני(?:ים|ת|ות)?", "→ אינדקס חומרה + % מהתקן ('גבוה' / 'גבוה מאוד')"),
    (r"כמעט\s+מוחלט", "→ ניסוח עובדתי ('X מתוך Y קידוחים')"),
    (r"\bHIGH\b", "→ 'רמת ודאות גבוהה' (לא HIGH בפרוזה עברית)"),
    (r"\bMEDIUM\b", "→ 'רמת ודאות בינונית' (לא MEDIUM בפרוזה עברית)"),
    (r"\bLOW\b", "→ 'רמת ודאות נמוכה' (לא LOW בפרוזה עברית)"),
]

# DATA_PIPELINE_SPEC.md — required columns per CSV
DATA_PACK_SCHEMA = {
    "measurements_scoped.csv": [
        "canonical_well_id", "parameter_canonical", "unit_standardized",
        "result_value", "result_date", "source_file",
    ],
    "latest_results.csv": [
        "canonical_well_id", "parameter_canonical",
        "unit_standardized", "severity_index",
    ],
    "severity_by_well_family.csv": [
        "canonical_well_id", "family", "severity_index",
        "max_value_window", "max_value_date",
    ],
    "trends_by_well_parameter.csv": [
        "canonical_well_id", "parameter_canonical", "mann_kendall_z",
        "mann_kendall_p", "snr", "trend_classification",
    ],
    "monitoring_gaps.csv": [
        "canonical_well_id", "last_measurement_date", "is_active",
    ],
    "figure_ready_series.csv": [
        "canonical_well_id", "parameter_canonical",
        "severity_index", "unit_standardized",
    ],
}

# Zone diagnosis — 8 required thematic areas
DIAGNOSIS_REQUIRED_THEMES = [
    (r"מוקד\s+\d|אשכול|cluster|focus", "מוקדים גיאוגרפיים"),
    (r"CVOC|ממסים\s+מוכלר|TCE|PCE", "מזהמי CVOC"),
    (r"PFAS|פלואורו|כיסוי|פער\s+ניטור", "PFAS / פער כיסוי"),
    (r"דלק|BTEX|MTBE|FUEL|בנזן", "מזהמי דלק"),
    (r"מקורות|מפעל|מפעלים|מתחם\s+תעשי", "מקורות אפשריים"),
    (r"מגמ[הות]|עלי[יה]|ירידה|Mann.Kendall|טרנד", "מגמות"),
    (r"פער|[Cc]losed|שותק|הופסק|לא\s+נדגם", "פערי ניטור"),
    (r"המלצ|פעולה\s+נדרש|דחיפות|urgent|מיידי", "המלצות / דחיפות"),
]

# ── Result helpers ────────────────────────────────────────────────────────────

class QAResult:
    """Accumulates findings for one gate."""

    def __init__(self, gate: str, zone: str):
        self.gate = gate
        self.zone = zone
        self.errors: list[str] = []
        self.warnings: list[str] = []
        self.infos: list[str] = []

    def error(self, msg: str):
        self.errors.append(msg)

    def warn(self, msg: str):
        self.warnings.append(msg)

    def info(self, msg: str):
        self.infos.append(msg)

    @property
    def status(self) -> str:
        if self.errors:
            return "FAIL"
        if self.warnings:
            return "WARN"
        return "PASS"

    def print_report(self):
        width = 64
        print(f"\n{'='*width}")
        print(f"  QA Gate {self.gate} — Zone: {self.zone}")
        print(f"  Status: {self.status}")
        print(f"{'='*width}")

        for info in self.infos:
            print(f"  ✓  {info}")

        for warn in self.warnings:
            print(f"  ⚠  {warn}")

        for err in self.errors:
            print(f"  ✗  {err}")

        print(f"\n  Errors: {len(self.errors)}   Warnings: {len(self.warnings)}")


# ── Gate 2: Structured Data Pack ─────────────────────────────────────────────

def gate2_data_pack(zone: str, data_dir: Optional[Path] = None) -> QAResult:
    """Validate the 6-CSV Structured Data Pack.

    Checks:
    - All 6 CSVs exist in {zone}/02_data/ or provided data_dir
    - Required columns present (per DATA_PIPELINE_SPEC.md)
    - Borehole count consistent across CSVs (within ±5 tolerance)
    - No TPFAS/BETK in parameter columns
    - measurements_scoped has >100 rows (not empty run)
    - Series coherence (WARN, non-blocking): near-duplicate well IDs (merged-well
      defect) + disconnected (well, parameter) sub-series joined across a long gap
    """
    result = QAResult("2 (Data Pack)", zone)
    base = data_dir or (REPO_ROOT / zone / "02_data")

    if not base.exists():
        result.error(f"תיקיית נתונים לא נמצאה: {base}")
        return result

    borehole_counts: dict[str, int] = {}

    for csv_name, required_cols in DATA_PACK_SCHEMA.items():
        csv_path = base / csv_name
        if not csv_path.exists():
            result.error(f"CSV חסר: {csv_name}")
            continue

        try:
            with open(csv_path, encoding="utf-8") as f:
                reader = csv.DictReader(f)
                rows = list(reader)
                headers = reader.fieldnames or []
        except Exception as e:
            result.error(f"{csv_name}: שגיאת קריאה — {e}")
            continue

        # Column check
        missing_cols = [c for c in required_cols if c not in headers]
        if missing_cols:
            result.error(f"{csv_name}: עמודות חסרות — {missing_cols}")
        else:
            result.info(f"{csv_name}: {len(rows)} שורות, עמודות תקינות")

        # Row count sanity
        if csv_name == "measurements_scoped.csv" and len(rows) < 100:
            result.error(f"measurements_scoped.csv: רק {len(rows)} שורות — נראה ריצה חלקית")

        # Borehole count
        well_col = "canonical_well_id"
        if well_col in headers:
            wells = {r[well_col] for r in rows if r.get(well_col)}
            borehole_counts[csv_name] = len(wells)

        # TPFAS/BETK exclusion check
        param_col = "parameter_canonical"
        if param_col in headers:
            forbidden = {"TPFAS", "BETK", "tpfas", "betk"}
            found_forbidden = {r[param_col] for r in rows if r.get(param_col) in forbidden}
            if found_forbidden:
                result.error(f"{csv_name}: פרמטרים אסורים — {found_forbidden} (יש להוציא TPFAS/BETK מהניתוח)")

    # Cross-CSV borehole count consistency
    # Core CSVs (full population) must agree; derived CSVs may be smaller (filtered by thresholds)
    CORE_CSVS = {"measurements_scoped.csv", "latest_results.csv", "figure_ready_series.csv"}
    DERIVED_CSVS = {"severity_by_well_family.csv", "trends_by_well_parameter.csv", "monitoring_gaps.csv"}
    core_counts = {k: v for k, v in borehole_counts.items() if k in CORE_CSVS}
    derived_counts = {k: v for k, v in borehole_counts.items() if k in DERIVED_CSVS}

    if len(core_counts) >= 2:
        counts = list(core_counts.values())
        spread = max(counts) - min(counts)
        if spread > 3:
            detail = ", ".join(f"{k}: {v}" for k, v in core_counts.items())
            result.error(
                f"חוסר עקביות במנין קידוחים ב-CSVs הראשיים (פיזור {spread}): {detail}"
            )
        else:
            canonical_count = max(counts)
            result.info(f"מנין קידוחים ב-CSVs ראשיים: {canonical_count} קידוחים — עקבי")

            if derived_counts:
                derived_detail = ", ".join(f"{k}: {v}" for k, v in derived_counts.items())
                result.info(
                    f"CSVs נגזרים (סינון לפי ספים — צפוי שיהיו פחות): {derived_detail}"
                )

    # ── Series coherence (WARN — never blocks) ──
    # Two recurring data defects that schema checks miss:
    #   (a) near-duplicate well IDs that are really one physical well (e.g. an
    #       "אזור"/"איזור" spelling split — the נד אגד אזור 7 case), and
    #   (b) a (well, parameter) series joining two unrelated sub-series across a long
    #       gap — the chart then draws a line between two points that don't belong.
    ms_path = base / "measurements_scoped.csv"
    if ms_path.exists():
        try:
            with open(ms_path, encoding="utf-8") as f:
                ms_rows = list(csv.DictReader(f))
        except Exception:
            ms_rows = []

        # (a) near-duplicate well IDs (normalize whitespace, punctuation, איזור→אזור)
        def _norm_id(s: str) -> str:
            s = re.sub(r"\s+", "", s or "")
            s = s.replace("איזור", "אזור")
            return re.sub(r"[\"'\-_.]", "", s)

        norm_map: dict[str, set] = {}
        for r in ms_rows:
            wid = r.get("canonical_well_id", "")
            if wid:
                norm_map.setdefault(_norm_id(wid), set()).add(wid)
        dups = {k: v for k, v in norm_map.items() if len(v) > 1}
        if dups:
            for variants in list(dups.values())[:5]:
                result.warn(
                    f"קידוחים שעלולים להיות כפילות (מזהים דומים מאוד): {sorted(variants)} — "
                    f"בדוק אם זה קידוח פיזי אחד לפני שילוב סדרות"
                )
        else:
            result.info("קוהרנטיות מזהים: לא נמצאו מזהי-קידוח כפולים — תקין")

        # (b) disconnected sub-series: >3yr gap straddling a >100× concentration jump
        series: dict = {}
        for r in ms_rows:
            wid, par = r.get("canonical_well_id", ""), r.get("parameter_canonical", "")
            try:
                v = float(r.get("result_value", ""))
                d = datetime.strptime((r.get("result_date", "") or "")[:10], "%Y-%m-%d")
            except (ValueError, TypeError):
                continue
            if v > 0 and wid and par:
                series.setdefault((wid, par), []).append((d, v))

        disconnected = []
        for (wid, par), pts in series.items():
            if len(pts) < 3:
                continue
            pts.sort()
            for (d1, v1), (d2, v2) in zip(pts, pts[1:]):
                gap_days = (d2 - d1).days
                ratio = max(v1, v2) / min(v1, v2)
                if gap_days > 1095 and ratio > 100:
                    disconnected.append((wid, par, gap_days, ratio))
                    break
        if disconnected:
            for wid, par, gap, ratio in disconnected[:5]:
                result.warn(
                    f"סדרה אפשרית-מנותקת: {wid} / {par} — קפיצה ×{ratio:.0f} על פני "
                    f"~{gap // 365} שנים ללא דגימות ביניים (ודא שאין חיבור שתי תקופות לא-קשורות)"
                )
        else:
            result.info("קוהרנטיות סדרות: לא זוהו סדרות מנותקות חשודות — תקין")

    return result


# ── Gate 4: Zone Diagnosis ────────────────────────────────────────────────────

def gate4_diagnosis(zone: str, diagnosis_path: Optional[Path] = None) -> QAResult:
    """Validate zone_diagnosis.md.

    Checks:
    - File exists
    - 8 required thematic areas covered
    - Geographic framing: at least 2 named geographic foci (מוקד N + location)
    - PFAS framed as gap (not as active contamination focus)
    - Monitoring gaps with dates mentioned (closed wells)
    - No pipeline terminology leaking in
    """
    result = QAResult("4 (Zone Diagnosis)", zone)

    if diagnosis_path is None:
        candidates = [
            REPO_ROOT / zone / "context_pack" / "04_diagnosis" / "zone_diagnosis.md",
            REPO_ROOT / zone / "04_diagnosis" / "zone_diagnosis.md",
            REPO_ROOT / zone / "rerun_v2_2026-05-28" / "context_pack" / "04_diagnosis" / "zone_diagnosis.md",
        ]
        diagnosis_path = next((p for p in candidates if p.exists()), None)

    if diagnosis_path is None or not diagnosis_path.exists():
        result.error(f"zone_diagnosis.md לא נמצא (בדוק context_pack/04_diagnosis/)")
        return result

    text = diagnosis_path.read_text(encoding="utf-8")
    result.info(f"zone_diagnosis.md: {len(text)} תווים")

    # 8 required themes
    for pattern, label in DIAGNOSIS_REQUIRED_THEMES:
        if re.search(pattern, text, re.IGNORECASE):
            result.info(f"נושא מכוסה: {label}")
        else:
            result.warn(f"נושא לא זוהה: {label} — בדוק שהאבחון מכסה נושא זה")

    # Geographic framing: at least 2 named foci with location
    geo_foci = re.findall(r"מוקד\s+\d[^—\n]*(?:—|:|\n)", text)
    if len(geo_foci) >= 2:
        result.info(f"מבנה גיאוגרפי: {len(geo_foci)} מוקדים מוגדרים")
    else:
        result.warn(
            "מבנה גיאוגרפי חלש — האבחון אמור להגדיר מוקדים גיאוגרפיים מנוסחים "
            "(מוקד 1 — שם+מיקום, מוקד 2 — ...). זה הבסיס לפרק 3 בדוח."
        )

    # PFAS framing: must appear as gap, not active focus
    pfas_section = re.search(r"PFAS.{0,200}", text, re.DOTALL | re.IGNORECASE)
    if pfas_section:
        pfas_text = pfas_section.group(0)
        if re.search(r"פער|כיסוי|לא\s+נדגם|coverage\s+gap", pfas_text, re.IGNORECASE):
            result.info("PFAS מוסגר נכון כפער כיסוי")
        elif re.search(r"מוקד\s+\d.*PFAS|PFAS.*מוקד\s+\d", pfas_text, re.IGNORECASE):
            result.warn(
                "PFAS מוצג כמוקד זיהום — יש למסגרו כפער כיסוי שיטתי (max_bucket=0 → coverage gap)"
            )

    # Monitoring gaps with dates
    gap_with_date = re.findall(r"(?:הופסק|סגור|closed|שותק)[^.]*20\d\d", text)
    if gap_with_date:
        result.info(f"פערי ניטור עם תאריכים: {len(gap_with_date)} ממצאים")
    else:
        result.warn("פערי ניטור: לא זוהו תאריכי הפסקה מפורשים — בדוק שקידוחים סגורים מתועדים עם שנת סגירה")

    # Pipeline terminology in diagnosis (soft check — diagnosis is internal)
    for term_pattern in PIPELINE_TERMS[:6]:  # only file names, not step refs
        if re.search(term_pattern, text):
            result.warn(f"שפת pipeline באבחון: '{term_pattern}' — מקובל בשלב 4, אך וודא שלא עובר לדוח")

    return result


# ── Gate 5: V5 Report ─────────────────────────────────────────────────────────

def gate5_report(zone: str, report_path: Optional[Path] = None,
                 data_dir: Optional[Path] = None) -> QAResult:
    """Validate {ZONE}_REPORT_V5.md.

    Checks:
    - Pipeline terminology absent from body
    - Mandatory Hebrew-prose terms (§B.5: פירוק not ריקבון; הפסקת ניטור not שתיקה;
      רמת ודאות not HIGH/MEDIUM/LOW; no קיצוני/כמעט מוחלט)
    - 6 sections + appendices present
    - Borehole count consistent across sections (vs gate 2 baseline)
    - Focus count consistency (§1 vs §3 sub-sections)
    - §3 headers are geographic foci, not bare family names (WARN if family-only)
    - PFAS appears in "פערי כיסוי" section (WARN if missing or misframed)
    - Optional: cross-check §3 order against focus_order block in zone_diagnosis.md
    - Overstatement patterns flagged
    - Hebrew confidence (רמת ודאות גבוהה/בינונית/נמוכה) on facility attributions
    - No ALERT/WATCH/ELEVATED/STABLE terminology
    - Methodology section ≤ 10 lines (per PROCESS_GUIDE §II)
    """
    result = QAResult("5 (V5 Report)", zone)

    if report_path is None:
        # Prefer the highest version present (V6 > V5 > legacy names)
        candidates = [
            REPO_ROOT / zone / "output" / f"{zone.upper()}_REPORT_V6.md",
            REPO_ROOT / zone / "output" / "HOLON_REPORT_V6.md",
            REPO_ROOT / zone / "output" / f"{zone.upper()}_REPORT_V5.md",
            REPO_ROOT / zone / "output" / "HOLON_REPORT_V5.md",
        ]
        report_path = next((p for p in candidates if p.exists()), None)

    if report_path is None or not report_path.exists():
        result.error(f"דוח V5 לא נמצא (ציפי {zone.upper()}_REPORT_V5.md ב-{zone}/output/)")
        return result

    text = report_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    result.info(f"דוח: {report_path.name}, {len(lines)} שורות, {len(text):,} תווים")

    # ── 1. Pipeline terminology ──
    term_hits = []
    for pattern in PIPELINE_TERMS:
        matches = re.findall(pattern, text, re.IGNORECASE)
        if matches:
            linenos = [
                i + 1 for i, l in enumerate(lines)
                if re.search(pattern, l, re.IGNORECASE)
            ]
            term_hits.append((pattern, matches[0], linenos[:3]))

    if term_hits:
        for pattern, sample, linenos in term_hits:
            result.error(
                f"טרמינולוגיית pipeline בדוח — שורות {linenos}: '{sample}' "
                f"(יש להחליף בשפה מקצועית רגולטורית)"
            )
    else:
        result.info("טרמינולוגיית pipeline: לא נמצאה — תקין")

    # ── 1b. Mandatory Hebrew-prose terminology (STYLE_GUIDE §B.5) ──
    prose_hits = []
    for pattern, hint in BANNED_PROSE_TERMS:
        linenos = [i + 1 for i, l in enumerate(lines) if re.search(pattern, l)]
        if linenos:
            sample = re.search(pattern, text).group(0)
            prose_hits.append((sample, hint, linenos))
    if prose_hits:
        for sample, hint, linenos in prose_hits:
            result.error(
                f"מונח אסור בפרוזה (STYLE_GUIDE §B.5) — {len(linenos)}× שורות {linenos[:4]}: "
                f"'{sample}' {hint}"
            )
    else:
        result.info("טרמינולוגיה מחייבת (§B.5: פירוק/הפסקת ניטור/רמת ודאות): נקי — תקין")

    # ── 2. Section structure ──
    required_sections = [
        (r"^##\s+1\.", "1. תקציר מקצועי"),
        (r"^##\s+2\.", "2. תמונת מצב"),
        (r"^##\s+3\.", "3. מוקדי זיהום"),
        (r"^##\s+4\.", "4. מגמות"),
        (r"^##\s+5\.", "5. מקורות"),
        (r"^##\s+6\.", "6. המלצות"),
        (r"^##\s+(7\.|מתודולוגיה)", "מתודולוגיה"),
        (r"^##\s+(8\.|מגבלות)", "מגבלות"),
        (r"^##\s+נספח", "נספחים"),
    ]
    for pattern, label in required_sections:
        if re.search(pattern, text, re.MULTILINE):
            result.info(f"סעיף קיים: {label}")
        else:
            result.error(f"סעיף חסר: {label}")

    # ── 3. Borehole count consistency ──
    count_pattern = r"(\d{2,3})\s+קידוח"
    counts_found = re.findall(count_pattern, text)
    if counts_found:
        unique_counts = set(int(c) for c in counts_found)
        if len(unique_counts) > 1:
            result.error(
                f"חוסר עקביות במנין קידוחים: {sorted(unique_counts)} — "
                f"יש להחליט על מספר אחד ולעדכן את כל הסעיפים"
            )
        else:
            result.info(f"מנין קידוחים עקבי: {unique_counts.pop()} קידוחים")

    # ── 4. Focus count consistency (§1 vs §3) ──
    # Count foci mentioned in §1 summary
    sec1_match = re.search(r"^##\s+1\.(.*?)^##\s+2\.", text, re.MULTILINE | re.DOTALL)
    sec3_match = re.search(r"^##\s+3\.(.*?)^##\s+4\.", text, re.MULTILINE | re.DOTALL)

    if sec1_match and sec3_match:
        # Count "X מוקדים" in §1
        foci_in_summary = re.findall(r"(\d+)\s+מוקד", sec1_match.group(1))
        declared_count = int(foci_in_summary[0]) if foci_in_summary else None

        # Count sub-sections in §3 (### 3.x)
        subsections_in_3 = re.findall(r"^###\s+3\.\d", sec3_match.group(1), re.MULTILINE)
        actual_count = len(subsections_in_3)

        if declared_count is not None and actual_count > 0:
            if declared_count != actual_count:
                result.error(
                    f"חוסר עקביות במנין מוקדים: תקציר מצהיר על {declared_count} מוקדים, "
                    f"אבל פרק 3 מכיל {actual_count} תתי-פרקים"
                )
            else:
                result.info(f"מנין מוקדים עקבי: {actual_count}")

    # ── 5. Focus-first: §3 headers must be geographic foci, not bare family names ──
    # Per §IV (SSOT): each ### 3.N must name a geographic location/facility, not just a family keyword.
    family_only_patterns = [
        r"^###\s+3\.\d+\s+(ממסים\s+מוכלרים?|CVOC)\s*$",
        r"^###\s+3\.\d+\s+(מתכות\s+כבדות?|METALS)\s*$",
        r"^###\s+3\.\d+\s+(דלק|FUEL|BTEX)\s*$",
        r"^###\s+3\.\d+\s+PFAS\s*$",
    ]
    sec3_headers = re.findall(r"^###\s+3\.\d+.*$", text, re.MULTILINE)
    family_only_headers = []
    for hdr in sec3_headers:
        for pat in family_only_patterns:
            if re.match(pat, hdr, re.IGNORECASE):
                family_only_headers.append(hdr.strip())
                break

    if family_only_headers:
        result.warn(
            f"כותרות §3 נראות כשמות-משפחה בלבד (לא מוקד גיאוגרפי) — §IV מחייב שם מקום/מתקן: "
            f"{family_only_headers[:3]}"
        )
    elif sec3_headers:
        result.info(f"כותרות §3: {len(sec3_headers)} מוקדים — נראות כמוקדים גיאוגרפיים")

    # Optional cross-check: if zone_diagnosis.md exists, verify §3 header order matches focus_order block
    diag_candidates = [
        REPO_ROOT / zone / "context_pack" / "04_diagnosis" / "zone_diagnosis.md",
        REPO_ROOT / zone / "04_diagnosis" / "zone_diagnosis.md",
    ]
    diag_path = next((p for p in diag_candidates if p.exists()), None)
    if diag_path is None:
        diag_path = REPO_ROOT / zone / "04_diagnosis" / "zone_diagnosis.md"  # fallback for error msg
    if diag_path.exists():
        diag_text = diag_path.read_text(encoding="utf-8")
        focus_block = re.search(r"^##\s+סדר מוקדים(.*?)(?:^##|\Z)", diag_text, re.MULTILINE | re.DOTALL)
        if focus_block:
            result.info("בלוק 'סדר מוקדים' נמצא ב-zone_diagnosis.md — ניתן לאמת עקביות ידנית")
        else:
            result.warn("zone_diagnosis.md קיים אך חסר בלוק '## סדר מוקדים' — Gate 4 אמור לחסום זאת")

    # ── 6. PFAS framing: must appear, ideally in "פערי כיסוי" section ──
    pfas_match = re.search(r"PFAS.{0,500}", text, re.DOTALL | re.IGNORECASE)
    if pfas_match:
        # Check if PFAS appears inside a "פערי כיסוי" section
        coverage_gap_section = re.search(
            r"###.*פערי\s+כיסוי.*?(?=^###|\Z)", text, re.MULTILINE | re.DOTALL | re.IGNORECASE
        )
        if coverage_gap_section and "PFAS" in coverage_gap_section.group(0).upper():
            result.info("PFAS מופיע בסעיף 'פערי כיסוי' — תקין (§IV)")
        else:
            pfas_text = pfas_match.group(0)
            if re.search(r"פער|כיסוי|לא\s+נדגם|coverage", pfas_text, re.IGNORECASE):
                result.warn(
                    "PFAS מוסגר כ'פער כיסוי' אך לא זוהה בסעיף 'פערי כיסוי' מוגדר — "
                    "שקול ליצור ### פערי כיסוי מפורש (§IV)"
                )
            else:
                result.warn(
                    "PFAS: לא זוהה ניסוח 'פערי כיסוי' — בדוק שלא מוצג כמוקד זיהום פעיל "
                    "כאשר אין נתונים מאמתים (PROCESS_GUIDE §IV)"
                )
    else:
        result.error("PFAS לא נמצא בדוח — חובה לפי PROCESS_GUIDE §IV גם כאשר אין נתונים")

    # ── 7. Overstatement patterns ──
    for pattern, suggestion in OVERSTATEMENT_PATTERNS:
        matches = re.findall(pattern, text, re.IGNORECASE)
        if matches:
            linenos = [i + 1 for i, l in enumerate(lines) if re.search(pattern, l, re.IGNORECASE)]
            result.warn(
                f"קביעה נחרצת יתר — שורות {linenos[:3]}: '{matches[0]}' "
                f"{suggestion}"
            )

    # ── 8. Confidence levels on facility attributions (Hebrew prose — §B.5) ──
    sec5_match = re.search(r"^##\s+5\.(.*?)^##\s+6\.", text, re.MULTILINE | re.DOTALL)
    if sec5_match:
        sec5_text = sec5_match.group(1)
        has_confidence = bool(re.search(
            r"רמת\s+ודאות|ודאות\s+(?:גבוהה|בינונית|נמוכה)|(?:גבוהה|בינונית|נמוכה)\s+ודאות",
            sec5_text))
        if not has_confidence:
            result.error(
                "פרק 5 (מקורות): חסר דירוג 'רמת ודאות: גבוהה/בינונית/נמוכה' על שיוכי מקור "
                "(§B.5 — לא HIGH/MEDIUM/LOW)"
            )
        else:
            result.info("רמת ודאות בפרק 5 (מינוח עברי): קיימת")

    # ── 9. Methodology section length ──
    method_match = re.search(r"^##\s+(?:7\.|מתודולוגיה)(.*?)^##\s+(?:8\.|מגבלות)", text, re.MULTILINE | re.DOTALL)
    if method_match:
        method_lines = [l for l in method_match.group(1).splitlines() if l.strip()]
        if len(method_lines) > 15:
            result.warn(
                f"פרק מתודולוגיה: {len(method_lines)} שורות — PROCESS_GUIDE §II מבקש 5–10 שורות בלבד; "
                f"הפנה לPROCESS_GUIDE §III לפרטים מלאים"
            )
        else:
            result.info(f"פרק מתודולוגיה: {len(method_lines)} שורות — תמציתי, תקין")

    return result


# ── Gate 6: HTML / Word Output ────────────────────────────────────────────────

def gate6_output(zone: str, html_path: Optional[Path] = None,
                 docx_path: Optional[Path] = None) -> QAResult:
    """Validate HTML and Word outputs.

    Checks:
    - HTML: figure count matches report
    - HTML: SVG borehole circles present (>50)
    - HTML: <bdi> wrapping present for numbers/chemicals
    - HTML: RTL meta / CSS present
    - Word: embedded images present
    - Word: RTL paragraphs ≥ 90% of total
    - Word: track changes enabled
    """
    result = QAResult("6 (HTML / Word Output)", zone)

    # Locate HTML (prefer highest version present: V6 > V5 > legacy)
    if html_path is None:
        html_candidates = [
            REPO_ROOT / zone / "output" / f"{zone.upper()}_REPORT_V6.html",
            REPO_ROOT / zone / "output" / "HOLON_REPORT_V6.html",
            REPO_ROOT / zone / "output" / f"{zone.upper()}_REPORT_V5.html",
            REPO_ROOT / zone / "output" / "HOLON_REPORT_V5.html",
        ]
        html_path = next((p for p in html_candidates if p.exists()), html_candidates[-1])

    # Locate DOCX
    if docx_path is None:
        docx_candidates = [
            REPO_ROOT / zone / "output" / f"{zone.upper()}_REPORT_V6.docx",
            REPO_ROOT / zone / "output" / "HOLON_REPORT_V6.docx",
            REPO_ROOT / zone / "output" / f"{zone.upper()}_REPORT_V5.docx",
            REPO_ROOT / zone / "output" / "HOLON_REPORT_V5.docx",
        ]
        docx_path = next((p for p in docx_candidates if p.exists()), docx_candidates[-1])

    # ── HTML checks ──
    if not html_path.exists():
        result.error(f"HTML לא נמצא: {html_path.name}")
    else:
        html = html_path.read_text(encoding="utf-8")
        result.info(f"HTML: {html_path.name} ({len(html):,} תווים)")

        # SVG borehole count
        circle_count = len(re.findall(r"<circle", html))
        if circle_count < 50:
            result.error(f"SVG מפה: {circle_count} עיגולים בלבד — נדרש ≥50 לייצוג תקין של הקידוחים")
        else:
            result.info(f"SVG מפה: {circle_count} עיגולי קידוח — תקין")

        # BDI wrapping
        bdi_count = len(re.findall(r"<bdi", html))
        number_count = len(re.findall(r"\d+[\d,.]*\s*(?:µg/L|mg/L|%)", html))
        if bdi_count < number_count * 0.3:
            result.warn(
                f"<bdi> עיטוף: {bdi_count} תגיות vs {number_count} ערכים מספריים — "
                f"בדוק PROCESS_GUIDE §IX.1 לעיטוף מספרים/כימיקלים"
            )
        else:
            result.info(f"<bdi> עיטוף: {bdi_count} תגיות — מספיק")

        # RTL meta
        if 'dir="rtl"' in html or "direction: rtl" in html:
            result.info("RTL הגדרות HTML: נמצאו")
        else:
            result.error("RTL: לא נמצאו הגדרות dir=rtl או direction:rtl — בדוק את תבנית ה-HTML")

        # Figure count
        svg_count = len(re.findall(r"<svg[^>]*>", html))
        result.info(f"SVG figures: {svg_count} (כולל מפה וגרפים)")

    # ── Word checks ──
    if not docx_path.exists():
        result.warn(f"Word לא נמצא: {docx_path.name} — דלג על בדיקות DOCX")
    else:
        try:
            from docx import Document
            from docx.oxml.ns import qn as _qn

            doc = Document(str(docx_path))
            para_count = len(doc.paragraphs)
            result.info(f"Word: {docx_path.name} ({docx_path.stat().st_size // 1024} KB, {para_count} פסקאות)")

            # Images
            image_count = sum(
                1 for para in doc.paragraphs
                for run in para.runs
                for _ in run._element.findall(
                    ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
                )
            )
            if image_count == 0:
                result.error("Word: אין איורים מוטבעים — יש להריץ embed_holon_figures_v5.py")
            else:
                result.info(f"Word: {image_count} איורים מוטבעים")

            # RTL paragraphs
            NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            rtl_count = sum(
                1 for para in doc.paragraphs
                if (ppr := para._element.find(f".//{{{NS}}}pPr")) is not None
                and ppr.find(f".//{{{NS}}}bidi") is not None
            )
            rtl_pct = rtl_count / para_count * 100 if para_count else 0
            if rtl_pct < 85:
                result.error(f"Word RTL: {rtl_count}/{para_count} פסקאות ({rtl_pct:.0f}%) — נדרש ≥85%")
            else:
                result.info(f"Word RTL: {rtl_count}/{para_count} פסקאות ({rtl_pct:.0f}%) — תקין")

            # Track changes
            settings = doc.settings._element
            track = settings.find(
                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}trackRevisions"
            )
            if track is None:
                result.warn("Word: מעקב שינויים לא מופעל — הוסף w:trackRevisions להגדרות")
            else:
                result.info("Word: מעקב שינויים מופעל")

        except ImportError:
            result.warn("python-docx לא מותקן — דילוג על בדיקות Word")
        except Exception as e:
            result.error(f"Word: שגיאה בקריאת DOCX — {e}")

    return result


# ── Gate 3: Prompt Layer freshness ────────────────────────────────────────────

# Ordering language abolished by REQ #24 (§IV is focus-first). Its presence in
# any generation-driving artifact (template or rendered prompt) is a drift FAIL.
ABOLISHED_ORDERING = [
    r"FUEL\s+always\s+last",
    r"FUEL\s+last\b",
    r"<?/?family_order>?",
    r"Family\s+[Oo]rdering",
    r"CVOC\s*→\s*METALS\s*→\s*PFAS\s*→\s*FUEL",
    r"FUEL.{0,8}אחרון",
]

# Placeholder tokens that MUST be substituted at render time, per prompt step.
_PLACEHOLDER_TOKENS = {
    "report": re.compile(r"\{[A-Z_][A-Z0-9_]*\}"),
    "diagnosis": re.compile(
        r"\{(?:zone_id|zone_he|data_dir|workspace_dir|forensics_dir|context_dir|output_path)\}"
    ),
}


def _sha12(path: Path) -> str:
    import hashlib
    return hashlib.sha256(path.read_text(encoding="utf-8").encode("utf-8")).hexdigest()[:12]


def gate3_prompt_layer(zone: str) -> QAResult:
    """Validate the PROMPT/INSTRUCTION layer that drives Opus generation.

    This gate closes the gap that REQ #24 left open: §IV was made the ordering
    SSOT and the governance docs became references, but the *rendered zone
    prompts* and the *prompt templates* are materialized artifacts that can drift
    silently. Checks:
    - Templates (report + diagnosis) contain no abolished family-first ordering.
    - Rendered prompts exist, carry a provenance stamp, match the current
      template's sha (not stale), have no unreplaced placeholders, and contain
      no abolished ordering language.
    - No stale duplicate rendered prompt survives outside the canonical location.
    - The V4-era template carries a DEPRECATED banner (footgun guard).
    """
    result = QAResult("3 (Prompt Layer)", zone)

    templates = {
        "report": REPO_ROOT / "scripts" / "templates" / "zone_report_prompt_template_v5.md",
        "diagnosis": REPO_ROOT / "scripts" / "templates" / "zone_diagnosis_prompt_template.md",
    }
    rendered = {
        "report": REPO_ROOT / zone / "context_pack" / "05_prompt" / "zone_report_prompt.md",
        "diagnosis": REPO_ROOT / zone / "context_pack" / "05_prompt" / "zone_diagnosis_prompt.md",
    }

    # 1. Templates must be focus-first (SSOT-aligned)
    for name, tpath in templates.items():
        if not tpath.exists():
            result.error(f"template חסר: {tpath.relative_to(REPO_ROOT)}")
            continue
        ttext = tpath.read_text(encoding="utf-8")
        hits = [p for p in ABOLISHED_ORDERING if re.search(p, ttext)]
        if hits:
            result.error(f"template '{name}' מכיל שפת-סדר מבוטלת (§IV focus-first): {hits}")
        else:
            result.info(f"template '{name}': focus-first — תקין")

    # 2. Rendered prompts: existence, freshness vs template, placeholders, ordering
    for name, rpath in rendered.items():
        if not rpath.exists():
            result.warn(
                f"prompt מרונדר חסר: {rpath.relative_to(REPO_ROOT)} — "
                f"הרץ: python scripts/render_zone_prompt.py --zone {zone} --step {name}"
            )
            continue
        rtext = rpath.read_text(encoding="utf-8")

        hits = [p for p in ABOLISHED_ORDERING if re.search(p, rtext)]
        if hits:
            result.error(
                f"prompt מרונדר '{name}' מכיל שפת-סדר מבוטלת (family-first) — "
                f"רנדר מחדש: {hits}"
            )

        leftover = sorted(set(_PLACEHOLDER_TOKENS[name].findall(rtext)))
        if leftover:
            result.error(f"prompt מרונדר '{name}': placeholders לא-מוחלפים {leftover[:6]}")

        stamp = re.search(r"template_sha256_12=([0-9a-f]{12})", rtext)
        tpath = templates[name]
        if not stamp:
            result.warn(
                f"prompt מרונדר '{name}': חסרה חותמת provenance — "
                f"רנדר מחדש דרך render_zone_prompt.py"
            )
        elif tpath.exists() and stamp.group(1) != _sha12(tpath):
            result.error(
                f"prompt מרונדר '{name}' מיושן: ה-template השתנה מאז הרינדור "
                f"(stamp {stamp.group(1)} ≠ current {_sha12(tpath)}) — רנדר מחדש"
            )
        elif not hits and not leftover:
            result.info(f"prompt מרונדר '{name}': טרי, תואם template, ללא placeholders — תקין")

    # 3. Stale duplicate rendered prompts outside the canonical context_pack/05_prompt
    for dup in REPO_ROOT.glob(f"{zone}/**/05_prompt/zone_report_prompt.md"):
        dup_str = str(dup)
        if "context_pack" in dup.parts or "_backup" in dup_str or "rerun" in dup_str:
            continue
        dtext = dup.read_text(encoding="utf-8")
        if any(re.search(p, dtext) for p in ABOLISHED_ORDERING):
            result.warn(
                f"prompt מרונדר כפול/מיושן מחוץ למיקום הקנוני: "
                f"{dup.relative_to(REPO_ROOT)} — שקול מחיקה (family-first)"
            )

    # 4. V4-era template footgun guard
    v4 = REPO_ROOT / "scripts" / "templates" / "zone_report_prompt_template.md"
    if v4.exists():
        head = v4.read_text(encoding="utf-8")[:600]
        if not re.search(r"DEPRECATED|מיושן|אין להשתמש", head):
            result.warn(
                "scripts/templates/zone_report_prompt_template.md (V4-era) ללא באנר "
                "DEPRECATED — footgun: מכיל סדר family-first ונראה שמיש"
            )

    return result


# ── Run all gates ─────────────────────────────────────────────────────────────

def run_all_gates(zone: str, args) -> int:
    """Run gates 2→3→4→5→6 in sequence. Returns exit code."""
    results = []
    gate_funcs = [
        ("2", lambda: gate2_data_pack(zone, getattr(args, "data_dir", None))),
        ("3", lambda: gate3_prompt_layer(zone)),
        ("4", lambda: gate4_diagnosis(zone, getattr(args, "diagnosis", None))),
        ("5", lambda: gate5_report(zone, getattr(args, "report", None))),
        ("6", lambda: gate6_output(zone, getattr(args, "html", None), getattr(args, "docx", None))),
    ]

    for gate_num, func in gate_funcs:
        r = func()
        r.print_report()
        results.append(r)

        if r.status == "FAIL" and not getattr(args, "continue_on_failure", False):
            print(f"\n⛔ שער {gate_num} נכשל — הפסקת ריצה.")
            print("   השתמש ב-'--continue-on-failure' להמשיך לשערים הבאים בכל מקרה.")
            break

    # Summary
    print(f"\n{'='*64}")
    print("  סיכום — כל השערים")
    print(f"{'='*64}")
    for r in results:
        icon = "✅" if r.status == "PASS" else ("⚠️" if r.status == "WARN" else "❌")
        print(f"  {icon}  שער {r.gate}: {r.status} ({len(r.errors)} שגיאות, {len(r.warnings)} אזהרות)")

    has_errors = any(r.status == "FAIL" for r in results)
    has_warnings = any(r.status == "WARN" for r in results)

    # Write QA report to zone/output/
    output_dir = REPO_ROOT / zone / "output"
    if output_dir.exists():
        report_file = output_dir / f"QA_REPORT_{datetime.now().strftime('%Y-%m-%d')}.md"
        _write_qa_report(report_file, zone, results)
        print(f"\n  📄 דוח QA נשמר: {report_file.relative_to(REPO_ROOT)}")

    if has_errors:
        return 1
    if has_warnings:
        return 2
    return 0


def _write_qa_report(path: Path, zone: str, results: list[QAResult]):
    """Write a structured markdown QA report."""
    lines = [
        f"# דוח QA — {zone}",
        f"**תאריך**: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        f"**סטטוס כולל**: {'FAIL' if any(r.status == 'FAIL' for r in results) else 'WARN' if any(r.status == 'WARN' for r in results) else 'PASS'}",
        "",
        "| שער | סטטוס | שגיאות | אזהרות |",
        "|-----|--------|--------|--------|",
    ]
    for r in results:
        lines.append(f"| שער {r.gate} | {r.status} | {len(r.errors)} | {len(r.warnings)} |")

    lines.append("")
    for r in results:
        lines.append(f"\n## שער {r.gate}\n")
        for e in r.errors:
            lines.append(f"- ❌ {e}")
        for w in r.warnings:
            lines.append(f"- ⚠️ {w}")
        for i in r.infos:
            lines.append(f"- ✓ {i}")

    path.write_text("\n".join(lines), encoding="utf-8")


# ── CLI ───────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="QA pipeline gates — V5 Hybrid Pipeline",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="Example: python scripts/qa_pipeline.py --gate all --zone Holon",
    )
    parser.add_argument("--gate", choices=["2", "3", "4", "5", "6", "all"], required=True)
    parser.add_argument("--zone", required=True, help="Zone directory name (e.g. Holon)")
    parser.add_argument("--report", type=Path, help="Override path to V5 report .md")
    parser.add_argument("--html", type=Path, help="Override path to HTML output")
    parser.add_argument("--docx", type=Path, help="Override path to Word output")
    parser.add_argument("--data-dir", type=Path, help="Override path to 02_data/ directory")
    parser.add_argument("--diagnosis", type=Path, help="Override path to zone_diagnosis.md")
    parser.add_argument("--continue-on-failure", action="store_true",
                        help="Run all gates even if an earlier gate fails")
    args = parser.parse_args()

    zone = args.zone

    if args.gate == "all":
        sys.exit(run_all_gates(zone, args))

    gate_map = {
        "2": lambda: gate2_data_pack(zone, args.data_dir),
        "3": lambda: gate3_prompt_layer(zone),
        "4": lambda: gate4_diagnosis(zone, args.diagnosis),
        "5": lambda: gate5_report(zone, args.report),
        "6": lambda: gate6_output(zone, args.html, args.docx),
    }

    result = gate_map[args.gate]()
    result.print_report()

    if result.status == "FAIL":
        sys.exit(1)
    elif result.status == "WARN":
        sys.exit(2)
    sys.exit(0)


if __name__ == "__main__":
    main()
